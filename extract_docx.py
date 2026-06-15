#!/usr/bin/env python3
"""Extract text from .docx files of team members."""
import os
try:
    from docx import Document
except ImportError:
    print("python-docx not installed")
    import sys
    sys.exit(1)

base = "/Users/zhengbo/Desktop/月报/1A 2026"
docx_files = [
    "陈嘉阳4月报.docx",
    "陈嘉阳5月月报.docx",
    "吴超4月月报.docx",
    "吴超5月份月报.docx",
    "晏军4月（月报）.txt",
    "晏军5月月报.docx",
    "张建强4月 工作月报.docx",
    "张建强5月 工作月报.docx",
]

for fname in docx_files:
    path = os.path.join(base, fname)
    if not os.path.exists(path):
        print(f"NOT FOUND: {fname}")
        continue
    
    print(f"\n{'='*60}")
    print(f"=== {fname} ===")
    print(f"{'='*60}")
    
    if fname.endswith('.txt'):
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            print(f.read()[:2000])
        continue
    
    try:
        doc = Document(path)
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                print(t[:200])
        
        # Also check tables
        for i, table in enumerate(doc.tables):
            print(f"\n[Table {i+1}: {len(table.rows)} rows]")
            for row in table.rows:
                cells = [cell.text.strip()[:30] for cell in row.cells]
                print(f"  {' | '.join(cells)}")
    except Exception as e:
        print(f"Error reading: {e}")
